# formatter/format_tables.py

from docx.shared import Pt, Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_tables(doc):
    pars = list(doc.paragraphs)
    for tbl in doc.tables:
        # повторять строку заголовков
        if tbl.rows:
            tbl.rows[0].heading = True

        # стиль ячеек
        for row in tbl.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height      = Cm(0.8)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
                        run.font.bold = False

        # подпись над таблицей: слева, без отступов и без точки
        try:
            idx = pars.index(tbl._tbl.getparent())
            if idx > 0:
                cap = pars[idx - 1]
                if cap.text.lower().startswith("таблица"):
                    fmt = cap.paragraph_format
                    fmt.left_indent       = Cm(0)
                    fmt.first_line_indent = Cm(0)
                    fmt.alignment         = WD_ALIGN_PARAGRAPH.LEFT
                    # убираем точку в конце подписи
                    cap.text = cap.text.rstrip('.').rstrip()
        except ValueError:
            pass
